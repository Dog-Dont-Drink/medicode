"""Writing polish endpoints (grammar proofreading, drafting helpers, etc.)."""

from __future__ import annotations

import io
import uuid
import zipfile
from datetime import datetime, timezone
from typing import Literal
from urllib.parse import quote
import xml.etree.ElementTree as ET

from fastapi import APIRouter, Depends, File, UploadFile
from fastapi.responses import Response
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.dependencies import get_current_user
from app.core.exceptions import BadRequest, Forbidden, NotFound
from app.db.database import get_db
from app.db.models.user import User
from app.db.models.writing_polish import GrammarEdit, WritingDocument, WritingDocumentVersion
from app.schemas.polish import (
    GrammarDocumentCreateRequest,
    GrammarDocumentDetailResponse,
    GrammarDocumentSummaryResponse,
    GrammarEditDecisionRequest,
    GrammarEditResponse,
    GrammarPolishRequest,
    GrammarPolishResponse,
    GrammarUploadParseResponse,
    GrammarVersionResponse,
)
from app.services.grammar_polish_service import polish_grammar
from app.services.resource_service import consume_user_resources, ensure_paid_feature_access, ensure_sufficient_resources


router = APIRouter(prefix="/polish", tags=["表达润色"])


def _ensure_paid_ai_ready(current_user: User) -> int:
    ensure_paid_feature_access(current_user, "ai_interpretation")
    return ensure_sufficient_resources(current_user, "ai_interpretation")


async def _get_document_for_user(document_id: uuid.UUID, current_user: User, db: AsyncSession) -> WritingDocument:
    result = await db.execute(select(WritingDocument).where(WritingDocument.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFound("文稿不存在")
    if doc.user_id != current_user.id:
        raise Forbidden("无权访问该文稿")
    return doc


def _version_response(version: WritingDocumentVersion) -> GrammarVersionResponse:
    return GrammarVersionResponse(
        id=str(version.id),
        version_no=int(version.version_no),
        source_module=version.source_module,
        settings=version.settings,
        model=version.model,
        llm_tokens_used=int(version.llm_tokens_used or 0),
        created_at=version.created_at,
    )


def _edit_response(edit: GrammarEdit) -> GrammarEditResponse:
    return GrammarEditResponse(
        id=str(edit.id),
        sentence_index=int(edit.sentence_index),
        original_text=edit.original_text or "",
        revised_text=edit.revised_text or "",
        edit_types=list(edit.edit_types or []),
        reasons=list(edit.reasons or []),
        confidence=edit.confidence,
        changed=bool(edit.changed),
        accepted=edit.accepted,
    )


@router.get("/grammar/documents", response_model=list[GrammarDocumentSummaryResponse])
async def list_grammar_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(WritingDocument)
        .where(WritingDocument.user_id == current_user.id)
        .order_by(WritingDocument.updated_at.desc(), WritingDocument.created_at.desc())
    )
    docs = result.scalars().all()
    return [
        GrammarDocumentSummaryResponse(
            id=str(doc.id),
            title=doc.title,
            text_type=doc.text_type,  # type: ignore[arg-type]
            section_type=doc.section_type,  # type: ignore[arg-type]
            updated_at=doc.updated_at,
            created_at=doc.created_at,
        )
        for doc in docs
    ]


@router.post("/grammar/documents", response_model=GrammarDocumentDetailResponse)
async def create_grammar_document(
    payload: GrammarDocumentCreateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = WritingDocument(
        user_id=current_user.id,
        title=payload.title,
        raw_text=payload.raw_text,
        text_type=payload.text_type,
        section_type=payload.section_type,
        updated_at=datetime.now(timezone.utc),
    )
    db.add(doc)
    await db.flush()
    return GrammarDocumentDetailResponse(
        id=str(doc.id),
        title=doc.title,
        raw_text=doc.raw_text,
        text_type=doc.text_type,  # type: ignore[arg-type]
        section_type=doc.section_type,  # type: ignore[arg-type]
        versions=[],
    )


@router.get("/grammar/documents/{document_id}", response_model=GrammarDocumentDetailResponse)
async def get_grammar_document(
    document_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_document_for_user(document_id, current_user, db)
    versions_result = await db.execute(
        select(WritingDocumentVersion)
        .where(WritingDocumentVersion.document_id == doc.id)
        .order_by(WritingDocumentVersion.version_no.desc())
    )
    versions = versions_result.scalars().all()
    return GrammarDocumentDetailResponse(
        id=str(doc.id),
        title=doc.title,
        raw_text=doc.raw_text,
        text_type=doc.text_type,  # type: ignore[arg-type]
        section_type=doc.section_type,  # type: ignore[arg-type]
        versions=[_version_response(v) for v in versions],
    )


@router.post("/grammar/documents/{document_id}/polish", response_model=GrammarPolishResponse)
async def run_grammar_polish(
    document_id: uuid.UUID,
    payload: GrammarPolishRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    charged_resources = _ensure_paid_ai_ready(current_user)
    doc = await _get_document_for_user(document_id, current_user, db)

    polish_result = await polish_grammar(
        raw_text=payload.raw_text,
        text_type=payload.text_type,
        section_type=payload.section_type,
        strength=payload.strength,
        protect_terms=payload.protect_terms,
        preserve_structure=payload.preserve_structure,
    )

    max_no_result = await db.execute(
        select(func.coalesce(func.max(WritingDocumentVersion.version_no), 0))
        .where(WritingDocumentVersion.document_id == doc.id)
    )
    next_version_no = int(max_no_result.scalar() or 0) + 1

    version = WritingDocumentVersion(
        document_id=doc.id,
        version_no=next_version_no,
        content=polish_result.revised_text,
        source_module="grammar",
        settings={
            "text_type": payload.text_type,
            "section_type": payload.section_type,
            "strength": payload.strength,
            "protect_terms": payload.protect_terms,
            "preserve_structure": payload.preserve_structure,
        },
        model=polish_result.model,
        llm_tokens_used=polish_result.llm_tokens_used,
    )
    db.add(version)
    await db.flush()

    edits: list[GrammarEdit] = []
    for item in polish_result.edits:
        edit = GrammarEdit(
            document_id=doc.id,
            version_id=version.id,
            sentence_index=item.index,
            suffix=item.suffix,
            original_text=item.original,
            revised_text=item.revised,
            changed=item.changed,
            edit_types=item.edit_types or [],
            reasons=item.reasons or [],
            confidence=item.confidence,
            accepted=True if item.changed else None,
        )
        edits.append(edit)
        db.add(edit)

    doc.raw_text = payload.raw_text
    doc.text_type = payload.text_type
    doc.section_type = payload.section_type
    doc.updated_at = datetime.now(timezone.utc)
    db.add(doc)

    remaining_balance = await consume_user_resources(
        db=db,
        user=current_user,
        operation="polish_grammar",
        billed_resources=charged_resources,
        actual_tokens=polish_result.llm_tokens_used,
    )

    summary = {
        "changed_sentences": sum(1 for e in edits if e.changed),
        "total_sentences": len(edits),
        "edit_type_counts": _count_edit_types(edits),
        "model": polish_result.model,
        "llm_tokens_used": polish_result.llm_tokens_used,
    }

    return GrammarPolishResponse(
        document_id=str(doc.id),
        version=_version_response(version),
        revised_text=polish_result.revised_text,
        edits=[_edit_response(e) for e in edits],
        summary=summary,
        charged_resources=charged_resources,
        charged_tokens=charged_resources,
        resource_balance=remaining_balance,
    )


def _count_edit_types(edits: list[GrammarEdit]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for edit in edits:
        for t in (edit.edit_types or []):
            key = str(t)
            if not key:
                continue
            counts[key] = counts.get(key, 0) + 1
    return counts


@router.get("/grammar/documents/{document_id}/versions/{version_no}", response_model=GrammarPolishResponse)
async def get_grammar_version(
    document_id: uuid.UUID,
    version_no: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_document_for_user(document_id, current_user, db)
    version_result = await db.execute(
        select(WritingDocumentVersion)
        .where(
            WritingDocumentVersion.document_id == doc.id,
            WritingDocumentVersion.version_no == version_no,
        )
    )
    version = version_result.scalar_one_or_none()
    if not version:
        raise NotFound("版本不存在")

    edits_result = await db.execute(
        select(GrammarEdit)
        .where(GrammarEdit.version_id == version.id)
        .order_by(GrammarEdit.sentence_index.asc())
    )
    edits = edits_result.scalars().all()
    revised_text = "".join((e.revised_text or e.original_text or "") + (e.suffix or "") for e in edits)
    summary = {
        "changed_sentences": sum(1 for e in edits if e.changed),
        "total_sentences": len(edits),
        "edit_type_counts": _count_edit_types(list(edits)),
        "model": version.model,
        "llm_tokens_used": int(version.llm_tokens_used or 0),
    }

    return GrammarPolishResponse(
        document_id=str(doc.id),
        version=_version_response(version),
        revised_text=revised_text,
        edits=[_edit_response(e) for e in edits],
        summary=summary,
        charged_resources=0,
        charged_tokens=0,
        resource_balance=int(current_user.token_balance or 0),
    )


@router.patch("/grammar/edits/{edit_id}", response_model=GrammarEditResponse)
async def decide_grammar_edit(
    edit_id: uuid.UUID,
    payload: GrammarEditDecisionRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(GrammarEdit).where(GrammarEdit.id == edit_id))
    edit = result.scalar_one_or_none()
    if not edit:
        raise NotFound("修改项不存在")

    doc = await _get_document_for_user(edit.document_id, current_user, db)
    _ = doc

    edit.accepted = bool(payload.accepted)
    db.add(edit)
    await db.flush()
    return _edit_response(edit)


@router.post("/grammar/parse-upload", response_model=GrammarUploadParseResponse)
async def parse_grammar_upload(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    _ = current_user
    filename = (file.filename or "upload").strip()
    content = await file.read()
    if not content:
        raise BadRequest("上传文件为空")

    if filename.lower().endswith(".txt"):
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("utf-8", errors="replace")
        return GrammarUploadParseResponse(filename=filename, text=text)

    if not filename.lower().endswith(".docx"):
        raise BadRequest("仅支持 .docx 或 .txt 文件")

    # Lightweight docx parsing (no extra dependency).
    text = _extract_docx_text(content)
    if not text.strip():
        raise BadRequest("未能从 docx 中解析到有效文本")
    return GrammarUploadParseResponse(filename=filename, text=text)


def _extract_docx_text(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            xml_bytes = zf.read("word/document.xml")
    except Exception as exc:
        raise BadRequest("docx 文件解析失败") from exc

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError as exc:
        raise BadRequest("docx XML 解析失败") from exc

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for para in root.findall(".//w:p", ns):
        texts = [t.text for t in para.findall(".//w:t", ns) if t.text]
        if texts:
            paragraphs.append("".join(texts))
    return "\n\n".join(paragraphs)


@router.post("/grammar/documents/{document_id}/export")
async def export_grammar_document(
    document_id: uuid.UUID,
    version_no: int,
    fmt: Literal["docx", "md", "txt"] = "docx",
    apply_acceptance: bool = True,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_document_for_user(document_id, current_user, db)

    version_result = await db.execute(
        select(WritingDocumentVersion)
        .where(
            WritingDocumentVersion.document_id == doc.id,
            WritingDocumentVersion.version_no == version_no,
        )
    )
    version = version_result.scalar_one_or_none()
    if not version:
        raise NotFound("版本不存在")

    edits_result = await db.execute(
        select(GrammarEdit)
        .where(GrammarEdit.version_id == version.id)
        .order_by(GrammarEdit.sentence_index.asc())
    )
    edits = edits_result.scalars().all()

    if apply_acceptance:
        final_text = "".join(
            ((e.revised_text if e.accepted is not False else e.original_text) or "") + (e.suffix or "")
            for e in edits
        )
    else:
        final_text = version.content or ""

    safe_title = (doc.title or "document").strip() or "document"
    if fmt == "txt":
        encoded_name = quote(f"{safe_title}.txt")
        return Response(
            content=final_text,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
        )

    if fmt == "md":
        encoded_name = quote(f"{safe_title}.md")
        return Response(
            content=final_text,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
        )

    # docx
    docx_bytes = _render_docx(final_text)
    encoded_name = quote(f"{safe_title}.docx")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )


def _render_docx(text: str) -> bytes:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as exc:
        raise BadRequest("缺少 python-docx 依赖，无法导出 Word；请安装 python-docx 后重试") from exc

    document = Document()
    for block in (text or "").replace("\r\n", "\n").split("\n"):
        if block.strip():
            document.add_paragraph(block)
        else:
            document.add_paragraph("")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
